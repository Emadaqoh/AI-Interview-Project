import os
import PyPDF2
import docx
from groq import Groq
import json
import os
from dotenv import load_dotenv

# تأكد من وضع مفتاحك هنا
load_dotenv() # قراءة ملف .env
api_key = os.getenv("GROQ_API_KEY") # استدعاء المفتاحclient = Groq(api_key=GROQ_API_KEY)

# --- هذه هي الدالة التي كانت مفقودة وتسببت في الخطأ ---
def extract_text(file, filename):
    text = ""
    try:
        if filename.endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text()
        elif filename.endswith('.docx'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            # إذا كان ملف نصي عادي
            text = file.read().decode('utf-8')
        return text
    except Exception as e:
        print(f"Error extracting text: {e}")
        return None

# دالة التحليل باستخدام Groq
def analyze_with_groq(cv_text):
    prompt = f"""
        Extract the following fields from the CV text into a valid JSON object. 
        If a field is missing, leave it as an empty string "" or empty list [].

        Fields:
        - personal_info: {{full_name, dob, nationality}}
        - contact_location: {{email, phone, website, location, availability}}
        - objectives: {{objective_text, summary_text, expected_salary}}
        - education_certs: [] (list of strings)
        - work_experience: [] (list of strings)
        - skills: []
        - languages: []

        CV Text: {cv_text[:7000]}
        """
    try:
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
            response_format={"type": "json_object"}
        )
        # تحويل النص المستلم إلى كائن بايثون
        return json.loads(chat_completion.choices[0].message.content)
    except Exception as e:
        print(f"Error: {e}")
        return {"full_name": "Error", "skills": [], "experience": "", "education": ""}



# import PyPDF2
# import docx
# import requests
# import json
# import io
# import re

# def extract_text(file, filename):
#     try:
#         if filename.endswith(".pdf"):
#             pdf_reader = PyPDF2.PdfReader(file)
#             text = ""
#             for page in pdf_reader.pages:
#                 text += page.extract_text() or ""
#             return text
#         elif filename.endswith(".docx"):
#             doc = docx.Document(file)
#             return "\n".join([p.text for p in doc.paragraphs])
#         return None
#     except Exception as e:
#         print(f"Extraction Error: {e}")
#         return ""

# def analyze_with_phi3(cv_text):
#     url = "http://localhost:11434/api/generate"
#     # رفع الحد لـ 2000 حرف لضمان الوصول لقسم الخبرة والتعليم
#     content_to_analyze = cv_text[:4000] 
    
#     prompt = f"""
#     Return ONLY a JSON object with these keys: 
#     "full_name", "skills" (list of strings), "experience" (string), "education" (string), "summary" (string).
#     CV Text: {content_to_analyze}
#     """
    
#     payload = {
#         "model": "phi3",
#         "prompt": prompt,
#         "stream": False,
#         "format": "json"
#     }
    
#     try:
#         response = requests.post(url, json=payload, timeout=200)
#         data = json.loads(response.json().get("response", "{}"))
        
#         # صمام أمان: التأكد من وجود المفاتيح المطلوبة لعدم كسر الرياكت
#         required_keys = ["full_name", "skills", "experience", "education", "summary"]
#         for key in required_keys:
#             if key not in data:
#                 data[key] = "Not found"
#         return data
        
#     except Exception as e:
#         print(f"Error: {e}")
#         return {"full_name": "Error", "skills": [], "experience": "N/A", "education": "N/A", "summary": "Error"}